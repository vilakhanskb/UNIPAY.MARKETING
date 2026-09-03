from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "documents" / "UniPay_Pitch_Transcript_LO.docx"
LOGO = ROOT / "assets" / "unipay-brand.png"

# Preset: launch_messaging_guide (compact_reference_guide alias).
# Named brand override: Noto Sans Lao replaces Calibri; UniPay ink/mint/gold replace preset blues.
FONT = "Noto Sans Lao"
INK = "082C32"
MINT = "079B91"
MINT_SOFT = "DFF8F4"
GOLD = "E4A62F"
GOLD_SOFT = "FFF4D8"
MUTED = "647B7F"
LINE = "DCE8E6"
PAPER = "F7FAF9"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for child_name in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(child_name))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.25, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def add_text(doc, text, size=11, color=INK, bold=False, italic=False, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, line=line)
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_kicker(doc, text):
    return add_text(doc, text.upper(), size=9, color=MINT, bold=True, after=4)


def add_callout(doc, label, text, fill=MINT_SOFT, accent=MINT):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=3, line=1.2)
    set_run_font(p.add_run(label + "  "), size=10, color=accent, bold=True)
    set_run_font(p.add_run(text), size=10, color=INK)
    add_text(doc, "", size=1, after=4)


def create_numbering(doc, fmt="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if fmt == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.25)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)
    set_run_font(p.add_run(text), size=11, color=INK)
    return p


def add_timeline_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["ເວລາ", "ຫົວຂໍ້", "ຈຸດປະສົງ"]
    for i, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], INK)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(value), size=9.5, color=WHITE, bold=True)
    for time, title, purpose in rows:
        cells = table.add_row().cells
        for i, value in enumerate((time, title, purpose)):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.15)
            set_run_font(p.add_run(value), size=9.5, color=INK, bold=(i == 1))
    apply_table_geometry(table, [1200, 2640, 5520])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    add_text(doc, "", size=1, after=4)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1)
    set_run_font(p.add_run("UniPay Pitch Transcript  |  ສະບັບສຳລັບ Demo"), size=8, color=MUTED)
    set_run_font(p.add_run("  •  "), size=8, color=GOLD)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    set_run_font(run, size=8, color=MUTED)
    run._r.extend([fld_begin, instr, fld_end])


def add_script_block(doc, timecode, title, show, say, transition=None):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    apply_table_geometry(table, [1500, 7860])
    # Keep each timed speaking block together so a transition or final line is
    # never stranded on the next page during Word/PDF rendering.
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    left, right = table.rows[0].cells
    set_cell_shading(left, INK)
    set_cell_shading(right, PAPER)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(timecode), size=10, color=WHITE, bold=True)
    p = right.paragraphs[0]
    style_paragraph(p, after=3, line=1.15)
    set_run_font(p.add_run(title), size=11.5, color=MINT, bold=True)
    p = right.add_paragraph()
    style_paragraph(p, after=3, line=1.2)
    set_run_font(p.add_run("ສະແດງ: "), size=9.5, color=GOLD, bold=True)
    set_run_font(p.add_run(show), size=9.5, color=INK)
    p = right.add_paragraph()
    style_paragraph(p, after=0, line=1.25)
    set_run_font(p.add_run("ຄຳເວົ້າ: "), size=9.5, color=MINT, bold=True)
    set_run_font(p.add_run(say), size=10.5, color=INK)
    if transition:
        p = right.add_paragraph()
        style_paragraph(p, before=3, after=0, line=1.2)
        set_run_font(p.add_run("ປະໂຫຍກເຊື່ອມ: "), size=9.5, color=MUTED, bold=True)
        set_run_font(p.add_run(transition), size=9.5, color=MUTED, italic=True)
    add_text(doc, "", size=1, after=2)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1.12)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Title": (29, INK, 0, 8),
        "Subtitle": (13.5, MUTED, 0, 18),
        "Heading 1": (16, MINT, 18, 10),
        "Heading 2": (13, MINT, 14, 7),
        "Heading 3": (12, INK, 10, 5),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.font.italic = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        p_pr = style._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is not None:
            p_pr.remove(p_bdr)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    bullet_id = create_numbering(doc, "bullet")
    number_id = create_numbering(doc, "decimal")

    # Cover: workshop_agenda pattern adapted to the UniPay brand.
    add_text(doc, "UNIPAY PLATFORM", size=10, color=MINT, bold=True, after=42)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo = p.add_run().add_picture(str(LOGO), width=Inches(0.72))
        logo._inline.docPr.set("descr", "UniPay Platform logo")
        logo._inline.docPr.set("title", "UniPay Platform")
        style_paragraph(p, after=22)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ບົດນຳສະເໜີ ແລະ Demo Script")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Mobile App · MiniApps · Backoffice Portal")
    add_callout(doc, "ເວລານຳສະເໜີ", "15–20 ນາທີ  |  ເນັ້ນຄຸນຄ່າທຸລະກິດ, Workflow ຕົວຈິງ ແລະການເລີ່ມ Pilot", fill=MINT_SOFT)
    add_text(doc, "ເອກະສານນີ້ຈັດທຳເປັນ Speaker Notes ສຳລັບຜູ້ນຳສະເໜີ. ສ່ວນ “ສະແດງ” ບອກສິ່ງທີ່ຄວນເປີດເທິງໜ້າຈໍ; ສ່ວນ “ຄຳເວົ້າ” ສາມາດອ່ານ ຫຼືປັບໃຫ້ເຂົ້າກັບຜູ້ຟັງໄດ້.", size=10.5, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "ເວັບປະກອບການນຳສະເໜີ", size=9, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "https://vilakhanskb.github.io/UNIPAY.MARKETING/", size=10, color=MINT, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_kicker(doc, "Run of show")
    add_heading(doc, "ລຳດັບການນຳສະເໜີ", 1)
    add_text(doc, "ໃຊ້ລຳດັບນີ້ເພື່ອໃຫ້ຜູ້ຟັງເຂົ້າໃຈບັນຫາ ກ່ອນເບິ່ງຟັງຊັນ ແລະຈົບດ້ວຍຂໍ້ສະເໜີ Pilot ທີ່ຕັດສິນໃຈໄດ້.")
    add_timeline_table(doc, [
        ("0:00–1:30", "ເປີດການນຳສະເໜີ", "ອະທິບາຍບັນຫາ, ກຸ່ມລູກຄ້າ ແລະຄຸນຄ່າຫຼັກ"),
        ("1:30–3:00", "UniPay Ecosystem", "ສະແດງການເຊື່ອມ Mobile App, MiniApps ແລະ Backoffice"),
        ("3:00–5:30", "Mobile App", "Login, KYC, ຕື່ມເງິນ LaoQR, ໂອນ, ຖອນ ແລະຊຳລະ QR"),
        ("5:30–10:30", "MiniApps", "UniPOS 3 Roles, UniMarket ແລະ UniSale"),
        ("10:30–15:30", "Backoffice", "ສະແດງ 6 Modules ແລະ Workflow ການກວດສອບ"),
        ("15:30–17:00", "Security & Governance", "SSO, Role, Scope, KYC, Audit, Monitoring ແລະ Runbook"),
        ("17:00–18:30", "KPI & Roadmap", "ກຳນົດ Baseline, Pilot metrics ແລະຂອບເຂດປັດຈຸບັນ"),
        ("18:30–20:00", "Decision Ask", "ຂໍຢືນຢັນຮ້ານ/ສາຂາ Pilot ແລະຂັ້ນຕອນຕໍ່ໄປ"),
    ])
    add_callout(doc, "ຈຸດສຳຄັນ", "ຖ້າມີເວລາພຽງ 15 ນາທີ ໃຫ້ຫຼຸດການເປີດລາຍລະອຽດ Sidebar ແຕ່ຮັກສາສ່ວນ Security, KPI ແລະ Decision Ask ໄວ້.", fill=GOLD_SOFT, accent=GOLD)

    doc.add_page_break()
    add_heading(doc, "ກຽມຄວາມພ້ອມກ່ອນ Demo", 1)
    prep = [
        "ເປີດ Website, Mobile App ແລະ Backoffice ໃຫ້ພ້ອມກ່ອນເລີ່ມ; ປິດ Notification ທີ່ບໍ່ກ່ຽວຂ້ອງ.",
        "ກວດ Network, ຂະໜາດຕົວໜັງສື ແລະ Screen sharing; ໃຊ້ຮູບ UI ໃນ Website ເປັນທາງສຳຮອງ.",
        "ກຽມ Demo accounts ໄວ້ນອກເອກະສານສາທາລະນະ. ຫ້າມສະແດງ Password, Token, OTP ຫຼືຂໍ້ມູນລູກຄ້າເທິງຈໍ.",
        "ເລືອກ Merchant ແລະສາຂາ Demo ລ່ວງໜ້າ ເພື່ອໃຫ້ Owner, Admin ແລະ Cashier ເຫັນຂອບເຂດທີ່ຖືກຕ້ອງ.",
        "ກຳນົດຂໍ້ມູນທົດລອງ: ສິນຄ້າ, ຈຳນວນ, ລາຄາ, Stock ແລະທຸລະກຳທີ່ສາມາດສະແດງຕໍ່ສາທາລະນະ.",
    ]
    for item in prep:
        add_list_item(doc, item, bullet_id)

    add_kicker(doc, "Full transcript")
    add_heading(doc, "1. ເປີດການນຳສະເໜີ", 1)
    add_script_block(doc, "0:00–0:40", "ທັກທາຍ ແລະວາງບໍລິບົດ", "ໜ້າ Hero ຂອງ Website", "ສະບາຍດີທຸກທ່ານ. ມື້ນີ້ພວກເຮົາຈະນຳສະເໜີ UniPay Platform ທີ່ເຊື່ອມການຂາຍ, ການສັ່ງຊື້, Wallet, ການຮັບຊຳລະ ແລະ Backoffice ເຂົ້າໄວ້ໃນ Workflow ດຽວ. ເປົ້າໝາຍບໍ່ແມ່ນພຽງການເພີ່ມອີກໜຶ່ງແອັບ, ແຕ່ແມ່ນການເຮັດໃຫ້ຂໍ້ມູນໜ້າຮ້ານ ແລະຫຼັງບ້ານກາຍເປັນຂໍ້ມູນຊຸດດຽວ.")
    add_script_block(doc, "0:40–1:30", "ບັນຫາ ແລະກຸ່ມລູກຄ້າ", "ເລື່ອນໄປ Executive Summary", "ໃນການດຳເນີນງານຈິງ Order, Stock, Payment ແລະລາຍງານມັກຢູ່ຫຼາຍບ່ອນ. ເຈົ້າຂອງຮ້ານເຫັນຂໍ້ມູນຊ້າ, ພະນັກງານຕ້ອງປ້ອນຊ້ຳ ແລະທີມ Support ໃຊ້ເວລາຫາສາເຫດ. UniPay ເໝາະຕັ້ງແຕ່ຮ້ານດຽວ, ທຸລະກິດຫຼາຍສາຂາ, Distributor ຈົນເຖິງ Partner ການຊຳລະ.", "ຕໍ່ໄປຈະເຫັນວ່າວຽກໃນ App ເຊື່ອມຫາ Backoffice ແນວໃດ.")

    add_heading(doc, "2. UniPay Ecosystem", 1)
    add_script_block(doc, "1:30–3:00", "ອະທິບາຍ App → Backoffice", "ສ່ວນ One connected workflow ແລະ Ecosystem", "ຝັ່ງຜູ້ໃຊ້ ແລະຮ້ານຄ້າຈະເຮັດວຽກຜ່ານ UniPay+ ແລະ MiniApps. ຝັ່ງທີມງານຈະໃຊ້ Backoffice ເພື່ອກຳນົດສິດ, ກວດ KYC, ຕິດຕາມ Order, Stock, Transaction, ການໄລ່ລຽງ ແລະ System Health. ສິ່ງສຳຄັນແມ່ນທຸກຂັ້ນຕອນມີສະຖານະ ແລະ Reference ສຳລັບກວດຄືນ.")

    add_heading(doc, "3. Mobile App", 1)
    add_script_block(doc, "3:00–3:40", "Login", "ເປີດໜ້າ Login; ບໍ່ສະແດງ Password", "ຜູ້ໃຊ້ສາມາດ Login ດ້ວຍ Email ແລະ Password. ຫຼັງຈາກອຸປະກອນໄດ້ຮັບອະນຸຍາດ ສາມາດໃຊ້ Biometric ເພື່ອເຂົ້າໃຊ້ໄດ້ໄວຂຶ້ນ. ເມນູ MiniApp ທີ່ເຫັນຈະຂຶ້ນກັບ Role ແລະຂອບເຂດທີ່ບັນຊີໄດ້ຮັບ.")
    add_script_block(doc, "3:40–4:15", "KYC", "ເປີດ Workflow KYC ຫຼືຮູບປະກອບ", "ກ່ອນໃຊ້ບໍລິການ Wallet ຕາມຂອບເຂດ ຜູ້ໃຊ້ສົ່ງຂໍ້ມູນ KYC ຜ່ານແອັບ. ຄຳຂໍຈະເຂົ້າຄິວໃນ Backoffice ເພື່ອໃຫ້ຜູ້ມີສິດກວດ, ບັນທຶກຜົນ ແລະເກັບ Audit evidence.")
    add_script_block(doc, "4:15–5:30", "Wallet ແລະ QR", "ເປີດໜ້າຕື່ມເງິນຜ່ານ LaoQR", "ໃນ Wallet ຜູ້ໃຊ້ສາມາດຕື່ມເງິນຜ່ານ LaoQR ໂດຍລະບຸຈຳນວນ, ກວດຄ່າທຳນຽມ ແລະຢືນຢັນຂໍ້ມູນກ່ອນສ້າງ QR. Workflow ດຽວກັນເຊື່ອມການໂອນ, ຖອນ ແລະຊຳລະ QR ເຂົ້າຫາ Transaction status ແລະ Ledger ໃນ Backoffice.", "ຈາກພື້ນຖານ Wallet ພວກເຮົາຈະເຂົ້າສູ່ MiniApps ສຳລັບທຸລະກິດ.")
    add_heading(doc, "4. MiniApps", 1)
    add_heading(doc, "4.1 UniPOS - 3 Roles", 2)
    add_script_block(doc, "5:30–6:25", "Owner", "ເປີດຮູບ Owner ໜ້າສັງລວມ/ລາຍງານ", "Owner ເຫັນຍອດຂາຍ, ຕົ້ນທຶນ, ກຳໄລ, ຈຳນວນ Order ແລະແນວໂນ້ມຕາມຊ່ວງເວລາ. ຄຸນຄ່າຂອງ Role ນີ້ແມ່ນການເຫັນພາບລວມໂດຍບໍ່ຕ້ອງລໍຖ້າລວບລວມລາຍງານຈາກແຕ່ລະສາຂາ.")
    add_script_block(doc, "6:25–7:20", "Branch Admin", "ເປີດຮູບ Admin ປັດຈຸບັນ", "Admin ຮັບຜິດຊອບວຽກປະຈຳຂອງສາຂາ: ຈັດການ Catalog, ເພີ່ມ ຫຼືແກ້ໄຂສິນຄ້າ, ກວດ Stock ແລະຕິດຕາມອໍເດີ. Owner ສາມາດມອບຂອບເຂດສາຂາໃຫ້ Admin ໂດຍບໍ່ຈຳເປັນມອບສິດທັງໝົດ.")
    add_script_block(doc, "7:20–8:35", "Cashier", "ເປີດໜ້າຂາຍ ແລະກົດເພີ່ມສິນຄ້າເຂົ້າກະຕ່າ", "Cashier ເຫັນສະເພາະຟັງຊັນໜ້າຮ້ານ. ພະນັກງານຄົ້ນຫາ ຫຼືສະແກນສິນຄ້າ, ເລືອກຈຳນວນ, ເພີ່ມລົງກະຕ່າ ແລະເລືອກຮັບຊຳລະດ້ວຍເງິນສົດ ຫຼື Lao QR. ການຂາຍເງິນສົດຍັງສາມາດເຂົ້າຄິວ Offline ແລະ Sync ເມື່ອ Network ກັບມາ.")
    add_callout(doc, "ຈຸດທີ່ຄວນເນັ້ນ", "3 Roles ໃຊ້ຂໍ້ມູນ Merchant/Branch ຊຸດດຽວກັນ ແຕ່ເຫັນເມນູ ແລະການດຳເນີນງານຕາມໜ້າວຽກ.")

    add_heading(doc, "4.2 UniMarket", 2)
    add_script_block(doc, "8:35–9:40", "Catalog ຫາ Order", "ເປີດ Catalog ແລະລາຍລະອຽດສິນຄ້າ", "UniMarket ເຊື່ອມຮ້ານຄ້າກັບ Catalog ແລະຜູ້ຈັດຈຳໜ່າຍ. ຜູ້ໃຊ້ເລືອກໝວດ, ຄົ້ນຫາສິນຄ້າ, ກວດ Variant, ລາຄາ, Promotion ແລະ Stock ກ່ອນເພີ່ມເຂົ້າກະຕ່າ. Backoffice ຮັບຊ່ວງຕໍ່ສຳລັບ Order status, Fulfilment, Reconciliation ແລະການໄລ່ລຽງ.")

    add_heading(doc, "4.3 UniSale", 2)
    add_script_block(doc, "9:40–10:30", "Sales network", "ເປີດໜ້າ UniSale ໃນ App", "UniSale ເປັນພື້ນທີ່ສຳລັບພະນັກງານຂາຍ ແລະເຄືອຂ່າຍຈັດຈຳໜ່າຍ ເພື່ອເບິ່ງຮ້ານ, ສິນຄ້າ ແລະຂອບເຂດທີ່ຮັບຜິດຊອບ. ໃນການນຳສະເໜີນີ້ ພວກເຮົາສະແດງພາບລວມການເຊື່ອມຕໍ່; ລາຍລະອຽດ Workflow ສາມາດນຳສະເໜີໃນຮອບຕໍ່ໄປ.")
    add_heading(doc, "5. Backoffice Portal", 1)
    add_text(doc, "ກ່ອນເປີດແຕ່ລະ Module ໃຫ້ເວົ້າວ່າ: “Sidebar ສະແດງສະເພາະ Function ທີ່ເປີດໃຊ້ປັດຈຸບັນ; Function ທີ່ບໍ່ເປີດໃຊ້ ຫຼືຊ້ຳກັນຈະບໍ່ນຳມາເປັນຂອບເຂດ Demo.”", color=MUTED, italic=True)
    add_script_block(doc, "10:30–11:40", "UniPOS Backoffice", "ເປີດ UniPOS Overview ແລະ Branch Stock", "Backoffice ຂອງ UniPOS ເຮັດໃຫ້ທີມງານເຫັນພາບລວມ, ຮ້ານ, ຄຳຂໍ Role, ພະນັກງານ, ອໍເດີ, ລູກຄ້າ, Catalog, ສາຂາ, Stock, ຫຼັກຖານນະໂຍບາຍ, Batch/Export ແລະສູນລາຍງານ. ຈຸດເດັ່ນແມ່ນເລືອກຂໍ້ມູນຕາມ Merchant, Branch ແລະຊ່ວງເວລາໄດ້.")
    add_script_block(doc, "11:40–12:35", "UniMarket Backoffice", "ເປີດ UniMarket Overview", "UniMarket ລວມຮ້ານ, Catalog, ສາຂາ, ອໍເດີ, ການກະທົບຍອດ, Auto Cancel, Support Diagnostics, Observability, Platform Reports, ໜ້າຮ້ານ, Priority Fees, Dealers, Salesmen, Stock ແລະ Reviews. ທຸກສະຖານະຊ່ວຍໃຫ້ຕິດຕາມ Order ຈາກການສັ່ງຈົນຮອດການຮັບສິນຄ້າ.")
    add_script_block(doc, "12:35–13:15", "ຈັດການຜູ້ໃຊ້", "ເປີດ User Management Overview", "Module ນີ້ມີພາບລວມ, ບົດຄວາມ, ຄຳເຫັນ ແລະການຈັດການສິດ. ການມອບ Role ແລະ Merchant scope ຊ່ວຍໃຫ້ຜູ້ໃຊ້ເຫັນສະເພາະວຽກທີ່ຮັບຜິດຊອບ.")
    add_script_block(doc, "13:15–14:00", "Wallet", "ເປີດ Wallet Overview ແລະຊີ້ Workflow LaoQR", "Wallet Backoffice ລວມພາບລວມ, KYC Review, Customer Investigation, ທຸລະກຳ, Affiliate Settings, Wallet Balance, ລາຍການໂອນລ່າສຸດ, Expense, Income ແລະ Wallet Audit. ເມື່ອລູກຄ້າຕື່ມເງິນຜ່ານ LaoQR ທີມງານສາມາດກວດ Payment reference, Callback, Wallet credit ແລະ Ledger ໄດ້.")
    add_script_block(doc, "14:00–14:45", "ລະບົບຮັບຊຳລະ", "ເປີດ Transaction Queue", "ທີມ Payment Operations ສາມາດຄົ້ນຫາທຸລະກຳ, ກວດ Partner, Webhook delivery, Payment Platform Health, Webhook SLA, Operations Signals ແລະຫ້ອງຄວບຄຸມການໄລ່ລຽງ. ລາຍງານແຍກ UniMarket, UniPOS ແລະ Partner ເຮັດໃຫ້ຮູ້ວ່າເງິນຢູ່ຂັ້ນຕອນໃດ.")
    add_script_block(doc, "14:45–15:30", "ຈັດການລະບົບ", "ເປີດ POS Dependency Status", "Platform Operations ຕິດຕາມ WebJob, Wallet Watch, Payment Jobs, Auto Cancel, Socket, Events, POS Dependency, POS Observability, Callback Investigation ແລະຂໍ້ມູນອ້າງອີງປະເພດຝາກ, ໂອນ ແລະຖອນ. ຈຸດປະສົງແມ່ນຮູ້ສັນຍານໄວ ແລະກວດຫາສາເຫດຈາກຂໍ້ມູນຊຸດດຽວ.")
    add_heading(doc, "6. Security & Governance", 1)
    add_script_block(doc, "15:30–17:00", "ສ້າງຄວາມເຊື່ອໝັ້ນ", "ເປີດສ່ວນ Security & Governance ໃນ Website", "UniPay ໃຊ້ Keycloak SSO ແລະ Role-based Access. ຂອບເຂດສາມາດກຳນົດຕາມ Merchant ແລະສາຂາ. ການສະແດງໃນ Portal ເປັນຊັ້ນນຳທາງ ແຕ່ Service API ແມ່ນຈຸດບັງຄັບສິດຫຼັກ. ລະບົບມີ KYC queue, ການປິດບັງຂໍ້ມູນ, Audit evidence, Monitoring, SLA ແລະ Runbook ສຳລັບ Operations. ໃນການ Demo ພວກເຮົາຈະບໍ່ສະແດງ Password, Token, Secret ຫຼືຂໍ້ມູນລູກຄ້າທີ່ລະບຸຕົວຕົນໄດ້.")

    add_heading(doc, "7. KPI, Roadmap ແລະ Pilot", 1)
    add_script_block(doc, "17:00–17:50", "Target KPI", "ເປີດ Pilot success metrics", "ພວກເຮົາບໍ່ຕັ້ງຕົວເລກຜົນສຳເລັດໂດຍບໍ່ມີຂໍ້ມູນ. ກ່ອນ Pilot ຈະບັນທຶກ Baseline ແລະກຳນົດ Target ຮ່ວມກັນ ໂດຍວັດ 4 ດ້ານ: ເວລາສະເລ່ຍຕໍ່ການຂາຍ, ອັດຕາ Stock ບໍ່ກົງ, ເວລາກວດຫາສາເຫດ ແລະຈຳນວນ/ອາຍຸຂອງລາຍການທີ່ຕ້ອງກວດໃນການໄລ່ລຽງ.")
    add_script_block(doc, "17:50–18:30", "Product status", "ເປີດ Roadmap", "ຂອບເຂດປັດຈຸບັນຂອງການນຳສະເໜີມີ 6 Modules. eCapital ແລະຄະແນນ–ລາງວັນເປັນ Roadmap ແລະບໍ່ນັບເປັນຄວາມສາມາດປັດຈຸບັນ. ການ Rollout ຈະເລີ່ມຈາກການເລືອກ Scope, ກຽມຂໍ້ມູນ/ສິດ, ທົດລອງ Go-live ແລະທົບທວນ KPI ກ່ອນຂະຫຍາຍ.")

    add_heading(doc, "8. ປິດການນຳສະເໜີ", 1)
    add_script_block(doc, "18:30–20:00", "Decision Ask", "ເປີດ CTA ທ້າຍ Website", "ສະຫຼຸບແລ້ວ UniPay ຊ່ວຍໃຫ້ວຽກໜ້າຮ້ານ, ການສັ່ງຊື້, Wallet, Payment ແລະ Operations ເຊື່ອມກັນດ້ວຍຂໍ້ມູນທີ່ກວດສອບໄດ້. ຂໍ້ສະເໜີຕໍ່ໄປແມ່ນໃຫ້ຢືນຢັນຮ້ານ ຫຼືສາຂາ Pilot, Workflow ທີ່ຈະທົດລອງ, ຜູ້ຮັບຜິດຊອບ ແລະ Target KPI. ຫຼັງຈາກນັ້ນທີມງານຈະຈັດເຮັດກຳນົດການ Onboarding ແລະ Go-live ຮ່ວມກັນ. ຂອບໃຈທຸກທ່ານ ແລະຍິນດີຕອບຄຳຖາມ.")
    add_callout(doc, "ຄຳຖາມສຳລັບປິດການປະຊຸມ", "“ທ່ານເຫັນວ່າຮ້ານ ຫຼື Workflow ໃດເໝາະທີ່ຈະເລີ່ມ Pilot ເປັນອັນດັບທຳອິດ?”", fill=GOLD_SOFT, accent=GOLD)
    doc.add_page_break()

    add_kicker(doc, "Demo operator guide")
    add_heading(doc, "ລຳດັບການກົດ Demo ແບບຫຍໍ້", 1)
    demo_steps = [
        "ເປີດ Website ໜ້າຫຼັກ → Executive Summary → App to Backoffice.",
        "ເປີດ Mobile App ທີ່ Login ໄວ້ແລ້ວ → ອະທິບາຍ Login/Biometric → KYC.",
        "Wallet → ຕື່ມເງິນ → LaoQR → ລະບຸຈຳນວນ → ກວດຂໍ້ມູນ. ຫ້າມຢືນຢັນທຸລະກຳຈິງໃນການ Demo.",
        "MiniApps → UniPOS → ສະແດງ Owner Overview → Admin Catalog/Stock → Cashier ເພີ່ມສິນຄ້າເຂົ້າກະຕ່າ.",
        "UniMarket → Catalog → Product detail → ອະທິບາຍ Order/Fulfilment; ບໍ່ສົ່ງ Order ຈິງ.",
        "UniSale → ສະແດງພາບລວມ ແລະຂອບເຂດພະນັກງານຂາຍ.",
        "Backoffice → UniPOS → UniMarket → Users → Wallet → Payments → System; ເປີດສະເພາະ 1–2 ໜ້າຫຼັກຕໍ່ Module.",
        "ກັບ Website → Security → KPI → Roadmap → Decision Ask.",
    ]
    for item in demo_steps:
        add_list_item(doc, item, number_id)

    add_heading(doc, "ແຜນສຳຮອງເມື່ອ Live Demo ຂັດຂ້ອງ", 1)
    fallback = [
        "ຖ້າ App ຫຼື Portal ຊ້າ: ໃຊ້ Real UI screenshots ໃນ Website ແລະອະທິບາຍ Workflow ຕາມ Transcript.",
        "ຖ້າ Login ບໍ່ສຳເລັດ: ບໍ່ພິມ Password ເທິງຈໍທີ່ Share; ຂ້າມໄປຮູບປະກອບ ແລະກັບມາຫຼັງຈາກຢຸດ Share.",
        "ຖ້າຂໍ້ມູນບາງລາຍການບໍ່ພ້ອມ: ບອກຕາມຄວາມຈິງວ່າກຳລັງຢູ່ໃນຂອບເຂດໃດ; ບໍ່ສ້າງຜົນລັບ ຫຼື KPI ຂຶ້ນເອງ.",
        "ຖ້າເວລາເຫຼືອໜ້ອຍ: ສະແດງ UniPOS 3 Roles, Wallet LaoQR, Payment Queue, Security ແລະ Decision Ask ເປັນຫຼັກ.",
    ]
    for item in fallback:
        add_list_item(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "ຂໍ້ຄວາມສັ້ນສຳລັບ Q&A", 1)
    qa_rows = [
        ("UniPay ແຕກຕ່າງແນວໃດ?", "ຈຸດເດັ່ນແມ່ນການເຊື່ອມ App, MiniApps ແລະ Backoffice ດ້ວຍ Role, Reference ແລະ Workflow ທີ່ກວດສອບໄດ້."),
        ("ຮອງຮັບຫຼາຍສາຂາບໍ?", "ຮອງຮັບ Merchant, ສາຂາ, Catalog, Stock ແລະສິດທີ່ແຍກຕາມຂອບເຂດ."),
        ("ຖ້າ Internet ຂາດ?", "UniPOS ຮອງຮັບການຂາຍເງິນສົດແບບ Offline queue; QR ແລະການດຳເນີນງານບາງປະເພດຕ້ອງ Online."),
        ("ຄວາມປອດໄພມີຫຍັງ?", "Keycloak SSO, Role/Scope, API enforcement, KYC, masked evidence, Audit ແລະ Monitoring/Runbook."),
        ("ເລີ່ມນຳໃຊ້ແນວໃດ?", "ເລີ່ມຈາກ Pilot ຂອບເຂດນ້ອຍ: ເລືອກຮ້ານ/ສາຂາ, ກຳນົດ Baseline, ຕັ້ງຄ່າ, ອົບຮົມ, Go-live ແລະທົບທວນ KPI."),
        ("eCapital ແລະ Rewards ພ້ອມບໍ?", "ບໍ່ຢູ່ໃນຂອບເຂດ Demo ປັດຈຸບັນ; ສະແດງເປັນ Roadmap ເທົ່ານັ້ນ."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for idx, text in enumerate(("ຄຳຖາມ", "ຄຳຕອບທີ່ແນະນຳ")):
        set_cell_shading(table.rows[0].cells[idx], INK)
        p = table.rows[0].cells[idx].paragraphs[0]
        set_run_font(p.add_run(text), size=9.5, color=WHITE, bold=True)
    for question, answer in qa_rows:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        style_paragraph(p, after=0, line=1.15)
        set_run_font(p.add_run(question), size=9.3, color=INK, bold=True)
        p = cells[1].paragraphs[0]
        style_paragraph(p, after=0, line=1.2)
        set_run_font(p.add_run(answer), size=9.3, color=INK)
    apply_table_geometry(table, [2600, 6760])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

    add_heading(doc, "ຂໍ້ມູນທີ່ຕ້ອງຢືນຢັນຫຼັງການນຳສະເໜີ", 1)
    decisions = [
        "ຮ້ານ/ສາຂາ ແລະ Workflow ທີ່ຈະໃຊ້ໃນ Pilot",
        "ຜູ້ຮັບຜິດຊອບຝັ່ງທຸລະກິດ, Operations ແລະ Technology",
        "ຄ່າ Baseline, Target KPI ແລະໄລຍະເວລາທົບທວນ",
        "ແຜນກຽມ Catalog/Stock, ການກຳນົດ Role ແລະການອົບຮົມ",
        "ວັນທີ Go-live, ຊ່ອງທາງ Support ແລະການປະຊຸມ Review",
    ]
    for item in decisions:
        add_list_item(doc, item, bullet_id)

    # Repeat footer/page settings if Word created extra sections.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1.12)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        add_footer(section)

    core = doc.core_properties
    core.title = "UniPay Pitch Transcript and Demo Script"
    core.subject = "Mobile App, MiniApps and Backoffice Portal presentation guide"
    core.author = "UniPay Platform"
    core.keywords = "UniPay, pitch, demo, transcript, Lao"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
